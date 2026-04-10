# utils/document_processor.py

import os
from datetime import datetime
import pdfplumber
from docx import Document as DocxDocument
from langchain_community.document_loaders import TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


def load_text_file(path: str) -> list[Document]:
    """
    Loads a .txt file line by line, tracking line numbers.
    Each document gets line_start and line_end metadata.
    """
    docs = []
    with open(path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # Group lines into ~paragraph-sized blocks (split on blank lines)
    current_block = []
    current_start = 1

    for i, line in enumerate(lines, start=1):
        if line.strip() == "":
            if current_block:
                docs.append(Document(
                    page_content="".join(current_block).strip(),
                    metadata={
                        "source": os.path.basename(path),
                        "line_start": current_start,
                        "line_end": i - 1,
                    }
                ))
                current_block = []
                current_start = i + 1
        else:
            current_block.append(line)

    # Last block
    if current_block:
        docs.append(Document(
            page_content="".join(current_block).strip(),
            metadata={
                "source": os.path.basename(path),
                "line_start": current_start,
                "line_end": len(lines),
            }
        ))

    return docs


def load_pdf_file(path: str) -> list[Document]:
    """
    Loads a PDF page by page.
    Each page becomes a separate Document with page_number metadata.
    """
    docs = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                docs.append(Document(
                    page_content=text.strip(),
                    metadata={
                        "source": os.path.basename(path),
                        "page_number": i,
                        "total_pages": len(pdf.pages),
                    }
                ))
    return docs


def load_docx_file(path: str) -> list[Document]:
    """
    Loads a .docx file paragraph by paragraph.
    Each paragraph gets a paragraph_index metadata.
    Empty paragraphs are skipped.
    """
    doc = DocxDocument(path)
    docs = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        docs.append(Document(
            page_content=text,
            metadata={
                "source": os.path.basename(path),
                "paragraph_index": i,
            }
        ))

    return docs


def load_documents_for_file(path: str) -> list[Document]:
    filename = os.path.basename(path)

    if filename.endswith(".txt"):
        return load_text_file(path)
    if filename.endswith(".pdf"):
        return load_pdf_file(path)
    if filename.endswith(".docx"):
        return load_docx_file(path)

    return []


def get_file_metadata(file_path: str) -> dict:
    file_stats = os.stat(file_path)
    return {
        'file_name': os.path.basename(file_path),
        'file_path': file_path,
        'file_size': file_stats.st_size,
        'modified_date': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
        'extension': os.path.splitext(file_path)[1]
    }


def split_documents_tiktoken(documents: list[Document], chunk_size=300, chunk_overlap=100) -> list[Document]:
    """
    Splits documents into chunks using tiktoken encoder.
    Preserves existing metadata (page_number, line_start, etc.)
    """
    text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    return text_splitter.split_documents(documents)


def add_chunk_metadata(docs_split: list[Document], file_path: str) -> list[Document]:
    """
    Adds chunk-level metadata to each split document.
    Preserves source-level metadata (page_number, line_start, etc.)
    """
    file_metadata = get_file_metadata(file_path)

    for i, doc in enumerate(docs_split):
        doc.metadata.update({
            "source": file_metadata['file_name'],
            "chunk_index": i,
            "total_chunks": len(docs_split),
            "chunk_size": len(doc.page_content),
            "file_size": file_metadata['file_size'],
            "file_path": file_metadata['file_path'],
            "upload_date": file_metadata['modified_date'],
            "extension": file_metadata['extension'],
        })

    return docs_split
